"""
Script d'extraction AMÉLIORÉ : PDF + Word + Sous-dossiers

Supporte :
- PDF (.pdf)
- Word moderne (.docx)  
- Word ancien (.doc) - avec avertissement
- Parcours récursif des sous-dossiers
"""

import os
from pathlib import Path
from pypdf import PdfReader
from docx import Document  # Pour .docx
import json
from datetime import datetime

# ============================================
# EXTRACTION PDF
# ============================================

def extract_text_from_pdf(pdf_path: str) -> dict:
    """
    Extrait le texte d'un PDF
    
    Comment ça marche :
    1. PdfReader lit le fichier PDF
    2. On boucle sur chaque page
    3. On extrait le texte de chaque page
    4. On combine tout
    """
    try:
        reader = PdfReader(pdf_path)
        
        text_parts = []
        for page_num, page in enumerate(reader.pages, start=1):
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
        
        full_text = "\n\n".join(text_parts)
        
        metadata = {
            "source": Path(pdf_path).name,
            "path": pdf_path,
            "format": "pdf",
            "num_pages": len(reader.pages),
            "num_chars": len(full_text),
            "extraction_date": datetime.now().isoformat()
        }
        
        return {
            "text": full_text,
            "metadata": metadata
        }
        
    except Exception as e:
        print(f"      ❌ Erreur PDF: {e}")
        return None


# ============================================
# EXTRACTION WORD (.docx)
# ============================================

def extract_text_from_docx(docx_path: str) -> dict:
    """
    Extrait le texte d'un fichier Word (.docx)
    
    Comment ça marche :
    1. Document() ouvre le .docx (c'est un ZIP avec du XML dedans)
    2. On boucle sur les paragraphes
    3. On récupère le texte de chaque paragraphe
    4. On combine avec double saut de ligne
    
    Note : .docx peut contenir tableaux, images, etc.
    On extrait uniquement le texte des paragraphes.
    """
    try:
        doc = Document(docx_path)
        
        # Extraire tous les paragraphes
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:  # Ignorer paragraphes vides
                paragraphs.append(text)
        
        # Combiner
        full_text = "\n\n".join(paragraphs)
        
        # Compter "pages" approximativement (1 page ≈ 500 mots)
        word_count = len(full_text.split())
        approx_pages = max(1, word_count // 500)
        
        metadata = {
            "source": Path(docx_path).name,
            "path": docx_path,
            "format": "docx",
            "num_pages": approx_pages,  # Approximatif
            "num_paragraphs": len(paragraphs),
            "num_chars": len(full_text),
            "extraction_date": datetime.now().isoformat()
        }
        
        return {
            "text": full_text,
            "metadata": metadata
        }
        
    except Exception as e:
        print(f"      ❌ Erreur DOCX: {e}")
        return None


# ============================================
# EXTRACTION WORD ANCIEN (.doc)
# ============================================

def extract_text_from_doc(doc_path: str) -> dict:
    """
    Tente d'extraire texte d'un .doc (ancien format)
    
    PROBLÈME : Format .doc est propriétaire Microsoft
    Solutions possibles :
    1. python-docx (parfois marche sur .doc)
    2. Conversion .doc → .docx (nécessite LibreOffice/Word)
    3. Bibliothèque spécialisée (antiword, mais complexe)
    
    Ici : On tente avec python-docx, sinon on avertit
    """
    print(f"      ⚠️  Format .doc ancien détecté")
    
    try:
        # Tenter quand même
        doc = Document(doc_path)
        
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
        
        full_text = "\n\n".join(paragraphs)
        
        if not full_text.strip():
            print(f"      ⚠️  Extraction vide - Convertir .doc → .docx recommandé")
            return None
        
        word_count = len(full_text.split())
        approx_pages = max(1, word_count // 500)
        
        metadata = {
            "source": Path(doc_path).name,
            "path": doc_path,
            "format": "doc",
            "num_pages": approx_pages,
            "num_paragraphs": len(paragraphs),
            "num_chars": len(full_text),
            "extraction_date": datetime.now().isoformat(),
            "warning": "Format .doc ancien - Qualité extraction incertaine"
        }
        
        return {
            "text": full_text,
            "metadata": metadata
        }
        
    except Exception as e:
        print(f"      ❌ Impossible de lire .doc: {e}")
        print(f"      💡 Conseil: Convertir '{Path(doc_path).name}' en .docx avec Word/LibreOffice")
        return None


# ============================================
# DISPATCHER : Choix du bon extracteur
# ============================================

def extract_text_from_file(file_path: str) -> dict:
    """
    Détecte le format et appelle le bon extracteur
    
    Supporte : .pdf, .docx, .doc
    """
    file_path = Path(file_path)
    extension = file_path.suffix.lower()
    
    if extension == ".pdf":
        return extract_text_from_pdf(str(file_path))
    elif extension == ".docx":
        return extract_text_from_docx(str(file_path))
    elif extension == ".doc":
        return extract_text_from_doc(str(file_path))
    else:
        print(f"      ⚠️  Format '{extension}' non supporté")
        return None


# ============================================
# PARCOURS RÉCURSIF
# ============================================

def extract_all_documents():
    """
    Parcourt RÉCURSIVEMENT data/raw/ et extrait tous les documents
    
    Nouveauté : utilise rglob() au lieu de glob()
    - glob("*.pdf") : seulement dossier actuel
    - rglob("*.pdf") : dossier actuel + sous-dossiers (RÉCURSIF)
    """
    
    raw_dir = Path("data/raw")
    categories = ["emploi_temps", "reglements", "procedures", "notes", "faqs"]
    
    all_documents = []
    stats = {
        "pdf": 0,
        "docx": 0,
        "doc": 0,
        "failed": 0
    }
    
    print("📚 EXTRACTION DES DOCUMENTS (PDF + Word)\n")
    print("=" * 60)
    
    for category in categories:
        category_path = raw_dir / category
        
        if not category_path.exists():
            continue
        
        # Trouver TOUS les fichiers (récursif)
        # rglob("*") = tous les fichiers dans tous les sous-dossiers
        all_files = list(category_path.rglob("*"))
        
        # Filtrer seulement les fichiers (pas les dossiers)
        document_files = [
            f for f in all_files 
            if f.is_file() and f.suffix.lower() in [".pdf", ".docx", ".doc"]
        ]
        
        if not document_files:
            print(f"\n📁 {category}: Aucun document trouvé")
            continue
        
        print(f"\n📁 {category}: {len(document_files)} document(s)")
        
        # Extraire chaque fichier
        for file_path in document_files:
            # Afficher le chemin relatif pour clarté
            relative_path = file_path.relative_to(category_path)
            print(f"   🔄 {relative_path}...")
            
            result = extract_text_from_file(str(file_path))
            
            if result:
                # Ajouter la catégorie
                result["metadata"]["category"] = category
                
                # Ajouter le chemin relatif (utile pour sous-dossiers)
                result["metadata"]["relative_path"] = str(relative_path)
                
                all_documents.append(result)
                
                # Stats
                format_type = result["metadata"]["format"]
                stats[format_type] += 1
                
                num_chars = result["metadata"]["num_chars"]
                num_pages = result["metadata"].get("num_pages", "?")
                print(f"      ✅ {num_pages} pages, {num_chars:,} caractères")
            else:
                stats["failed"] += 1
                print(f"      ❌ Échec extraction")
    
    print("\n" + "=" * 60)
    print(f"\n📊 Résumé extraction:")
    print(f"   - PDF extraits: {stats['pdf']}")
    print(f"   - DOCX extraits: {stats['docx']}")
    print(f"   - DOC extraits: {stats['doc']}")
    print(f"   - Échecs: {stats['failed']}")
    print(f"   - TOTAL RÉUSSI: {len(all_documents)}")
    
    return all_documents


# ============================================
# SAUVEGARDE
# ============================================

def save_extracted_documents(documents: list):
    """Sauvegarde en JSON"""
    
    output_dir = Path("data/processed")
    output_dir.mkdir(parents=True, exist_ok=True)
    
    output_file = output_dir / "extracted_documents.json"
    
    with open(output_file, "w", encoding="utf-8") as f:
        json.dump(documents, f, ensure_ascii=False, indent=2)
    
    print(f"\n💾 Sauvegardé dans: {output_file}")
    
    # Stats détaillées
    total_chars = sum(doc["metadata"]["num_chars"] for doc in documents)
    by_category = {}
    for doc in documents:
        cat = doc["metadata"]["category"]
        by_category[cat] = by_category.get(cat, 0) + 1
    
    print(f"\n📈 Statistiques finales:")
    print(f"   - Documents: {len(documents)}")
    print(f"   - Caractères totaux: {total_chars:,}")
    print(f"   - Moyenne par doc: {total_chars // len(documents):,}")
    print(f"\n   Par catégorie:")
    for cat, count in by_category.items():
        print(f"      • {cat}: {count}")


# ============================================
# MAIN
# ============================================

if __name__ == "__main__":
    print("\n🚀 Démarrage extraction...\n")
    
    documents = extract_all_documents()
    
    if documents:
        save_extracted_documents(documents)
        print("\n🎉 Extraction terminée avec succès !")
    else:
        print("\n⚠️ Aucun document extrait")